import sys
import glob
import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MEMBERS = ['강민채', '김민석', '오병욱']

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def add_member_heading(doc, name):
    p = doc.add_heading(name, level=1)
    for run in p.runs:
        run.font.name = '맑은 고딕'

def copy_body(src_doc, dst_doc):
    for elem in src_doc.element.body:
        if elem.tag.endswith('}sectPr'):
            continue
        dst_doc.element.body.append(deepcopy(elem))

def merge(week):
    out = Document()
    for p in list(out.paragraphs):
        p._element.getparent().remove(p._element)

    first = True
    for member in MEMBERS:
        pattern = os.path.join(member, week, '*.docx')
        files = glob.glob(pattern)
        if not files:
            print(f'[경고] {member}/{week}/ 에 docx 파일 없음')
            continue

        if not first:
            add_page_break(out)

        add_member_heading(out, member)
        copy_body(Document(files[0]), out)
        first = False

    os.makedirs('result', exist_ok=True)
    out_path = os.path.join('result', f'{week}.docx')
    out.save(out_path)
    print(f'완료: {out_path}')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('사용법: python result/merge_submissions.py week01')
        sys.exit(1)
    merge(sys.argv[1])
