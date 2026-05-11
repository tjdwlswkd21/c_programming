from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = '맑은 고딕'
    return p

def label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '맑은 고딕'
    run.font.size = Pt(11)
    return p

def blank_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

def code_block(doc):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run('// 소스코드를 여기에 붙여넣기')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    shading = p._element
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = shading.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def screenshot_placeholder(doc):
    p = doc.add_paragraph()
    run = p.add_run('[ 실행화면 캡처 이미지 삽입 ]')
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F9F9F9')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

# ── 표지 ──
heading(doc, 'Week 02 결과 보고서', 0)

info = doc.add_paragraph()
info.add_run('이름: ').bold = True
info.add_run('                    ')
info.add_run('  제출일: ').bold = True
info.add_run('              ')
for run in info.runs:
    run.font.name = '맑은 고딕'
    run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('─' * 60)
doc.add_paragraph()

# ── 문제별 섹션 ──
for i in range(1, 5):
    heading(doc, f'문제 {i}번', 1)

    label(doc, '▪ 문제 내용')
    doc.add_paragraph()
    doc.add_paragraph()
    blank_line(doc)

    label(doc, '▪ 소스코드')
    code_block(doc)
    code_block(doc)
    code_block(doc)
    blank_line(doc)

    label(doc, '▪ 실행화면')
    screenshot_placeholder(doc)
    blank_line(doc)

    label(doc, '▪ 피드백')

    p = doc.add_paragraph()
    r = p.add_run('어려움을 겪었던 부분')
    r.bold = True
    r.font.name = '맑은 고딕'
    r.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('문제를 통해 배우게 된 내용')
    r.bold = True
    r.font.name = '맑은 고딕'
    r.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph()

    if i < 4:
        doc.add_page_break()

doc.save('result/week02_제출양식.docx')
print('완료: result/week02_제출양식.docx')
